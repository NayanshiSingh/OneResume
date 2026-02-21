"""Export Service — handles file export (PDF/DOCX) and storage."""

import os
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings

logger = logging.getLogger(__name__)


def export_to_docx(resume_data: dict, output_path: str) -> str:
    """Export resume data to a DOCX file."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ── Header ────────────────────────────────────────────────
    pi = resume_data.get("personal_info", {})
    if pi:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(pi.get("full_name", ""))
        run.bold = True
        run.font.size = Pt(18)

        contact_parts = []
        if pi.get("email"):
            contact_parts.append(pi["email"])
        if pi.get("phone_number"):
            contact_parts.append(pi["phone_number"])

        if contact_parts:
            contact = doc.add_paragraph(" | ".join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Education ─────────────────────────────────────────────
    if resume_data.get("education"):
        doc.add_heading("Education", level=2)
        for edu in resume_data["education"]:
            text = f"{edu['degree']}"
            if edu.get("field_of_study"):
                text += f" in {edu['field_of_study']}"
            text += f" — {edu['institution']}"
            if edu.get("start_year") or edu.get("end_year"):
                years = f"{edu.get('start_year', '')} - {edu.get('end_year', '')}"
                text += f" ({years})"
            doc.add_paragraph(text)

    # ── Experience ────────────────────────────────────────────
    if resume_data.get("experience"):
        doc.add_heading("Experience", level=2)
        for exp in resume_data["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(exp["title"])
            run.bold = True
            if exp.get("subtitle"):
                p.add_run(f"  —  {exp['subtitle']}")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # ── Projects ──────────────────────────────────────────────
    if resume_data.get("projects"):
        doc.add_heading("Projects", level=2)
        for proj in resume_data["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(proj["title"])
            run.bold = True
            if proj.get("subtitle"):
                p.add_run(f"  |  {proj['subtitle']}")
            for bullet in proj.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # ── Skills ────────────────────────────────────────────────
    if resume_data.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(resume_data["skills"]))

    # ── Certifications ────────────────────────────────────────
    if resume_data.get("certifications"):
        doc.add_heading("Certifications", level=2)
        for cert in resume_data["certifications"]:
            text = cert["name"]
            if cert.get("issuing_organization"):
                text += f" — {cert['issuing_organization']}"
            if cert.get("year"):
                text += f" ({cert['year']})"
            doc.add_paragraph(text)

    # ── Achievements ──────────────────────────────────────────
    if resume_data.get("achievements"):
        doc.add_heading("Achievements", level=2)
        for ach in resume_data["achievements"]:
            text = ach["title"]
            if ach.get("description"):
                text += f": {ach['description']}"
            doc.add_paragraph(text, style="List Bullet")

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path
